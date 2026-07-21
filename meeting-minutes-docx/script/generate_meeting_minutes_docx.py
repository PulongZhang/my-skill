#!/usr/bin/env python3
"""Generate a meeting-minutes DOCX from the sanitized company template."""

from __future__ import annotations

import argparse
import json
import os
import tempfile
from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml.ns import qn


SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_TEMPLATE = SCRIPT_DIR.parent / "references" / "标准会议纪要模板.docx"
MEETING_TYPES = {
    "立项评审",
    "计划评审",
    "需求评审",
    "设计评审",
    "代码评审",
    "用例评审",
    "结项评审",
    "其他",
}
STATIC_PARTS = {
    "word/header1.xml",
    "word/_rels/header1.xml.rels",
    "word/media/image1.jpeg",
    "word/styles.xml",
    "word/theme/theme1.xml",
    "word/fontTable.xml",
    "word/settings.xml",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a styled meeting-minutes DOCX from JSON data."
    )
    parser.add_argument("--input", required=True, type=Path, help="Input JSON file")
    parser.add_argument("--output", required=True, type=Path, help="Output DOCX file")
    parser.add_argument(
        "--template",
        type=Path,
        default=DEFAULT_TEMPLATE,
        help=f"DOCX template (default: {DEFAULT_TEMPLATE})",
    )
    return parser.parse_args()


def replace_paragraph_text(paragraph, text: str) -> None:
    """Replace paragraph text while retaining the first run's formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = str(text)
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(str(text))


def replace_cell_text(cell, text: str) -> None:
    replace_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        replace_paragraph_text(paragraph, "")


def normalize_people(value) -> str:
    if isinstance(value, list):
        return "、".join(str(item).strip() for item in value if str(item).strip())
    return str(value or "").strip()


def resize_data_rows(table, count: int) -> None:
    """Keep one header row and exactly count styled data rows."""
    count = max(1, count)
    while len(table.rows) - 1 < count:
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    while len(table.rows) - 1 > count:
        table._tbl.remove(table.rows[-1]._tr)


def fill_rows(table, rows: list[list[str]]) -> None:
    resize_data_rows(table, len(rows))
    for row, values in zip(table.rows[1:], rows):
        for cell, value in zip(row.cells, values):
            replace_cell_text(cell, value)


def set_meeting_type(cell, selected: str) -> None:
    if selected not in MEETING_TYPES:
        raise ValueError(f"Unsupported meeting_type: {selected}")

    pending_symbol = None
    matched = False
    for run in cell.paragraphs[0]._p.findall(qn("w:r")):
        symbol = run.find(qn("w:sym"))
        text = run.find(qn("w:t"))
        if symbol is not None:
            symbol.set(qn("w:char"), "00A3")
            pending_symbol = symbol
        elif text is not None and pending_symbol is not None:
            label = (text.text or "").strip()
            if label == selected:
                pending_symbol.set(qn("w:char"), "0052")
                matched = True
            pending_symbol = None

    if not matched:
        raise ValueError(f"Meeting type checkbox not found: {selected}")


def set_footer_line(document: Document, recorder: str, meeting_date: str) -> None:
    paragraph = next(
        (
            item
            for item in document.paragraphs
            if "记录人：" in item.text and "日期：" in item.text
        ),
        document.paragraphs[-1],
    )
    replace_paragraph_text(
        paragraph, f"记录人：{recorder}\t日期：{meeting_date}"
    )
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.clear_all()
    section = document.sections[-1]
    usable_width = section.page_width - section.left_margin - section.right_margin
    tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)


def restore_static_parts(template: Path, output: Path) -> None:
    """Restore untouched style/header parts byte-for-byte after python-docx saves."""
    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=output.parent)
    os.close(fd)
    temp = Path(temp_name)
    try:
        with (
            ZipFile(template) as source_zip,
            ZipFile(output) as output_zip,
            ZipFile(temp, "w", ZIP_DEFLATED) as temp_zip,
        ):
            for info in output_zip.infolist():
                if info.filename in STATIC_PARTS and info.filename in source_zip.namelist():
                    content = source_zip.read(info.filename)
                else:
                    content = output_zip.read(info.filename)
                temp_zip.writestr(info, content)
        os.replace(temp, output)
    finally:
        if temp.exists():
            temp.unlink()


def validate_output(path: Path) -> None:
    with ZipFile(path) as archive:
        invalid_member = archive.testzip()
        if invalid_member:
            raise ValueError(f"Invalid DOCX member: {invalid_member}")

    document = Document(path)
    if not document.tables or len(document.tables[0].rows) != 8:
        raise ValueError("The generated DOCX lost the expected main table")
    footer = document.paragraphs[-1].text
    if "记录人：" not in footer or "日期：" not in footer or "\t" not in footer:
        raise ValueError("Recorder and date are not on the configured footer line")


def generate(data: dict, template: Path, output: Path) -> None:
    if not template.is_file():
        raise FileNotFoundError(f"Template not found: {template}")
    if output.suffix.lower() != ".docx":
        raise ValueError("Output filename must end with .docx")

    document = Document(template)
    main_table = document.tables[0]

    field_values = {
        (0, 1): data.get("meeting_topic", ""),
        (0, 3): data.get("project_number", ""),
        (1, 1): data.get("meeting_date", ""),
        (1, 3): data.get("meeting_format", ""),
        (2, 1): normalize_people(data.get("participants", "")),
        (2, 3): data.get("host", ""),
        (4, 1): data.get("conclusion", "待确认"),
    }
    for (row_index, cell_index), value in field_values.items():
        replace_cell_text(main_table.rows[row_index].cells[cell_index], value)

    set_meeting_type(main_table.rows[3].cells[1], data.get("meeting_type", "其他"))

    record_cell = main_table.rows[5].cells[0]
    replace_paragraph_text(record_cell.paragraphs[2], data.get("purpose", ""))
    issues = data.get("issues") or []
    issue_rows = [
        [
            str(index),
            item.get("questioner", ""),
            item.get("question", ""),
            item.get("owner", ""),
            item.get("solution", ""),
            item.get("result", "待确认"),
        ]
        for index, item in enumerate(issues, 1)
    ]
    if not issue_rows:
        issue_rows = [["1", "", "无", "", "", ""]]
    fill_rows(record_cell.tables[0], issue_rows)

    remaining_cell = main_table.rows[6].cells[0]
    remaining = data.get("remaining_issues") or []
    remaining_rows = [
        [
            str(index),
            item.get("description", ""),
            item.get("owner", "待确认"),
            item.get("due_date", "待确认"),
        ]
        for index, item in enumerate(remaining, 1)
    ]
    if not remaining_rows:
        remaining_rows = [["1", "无", "", ""]]
    fill_rows(remaining_cell.tables[0], remaining_rows)

    summary_cell = main_table.rows[7].cells[0]
    replace_paragraph_text(summary_cell.paragraphs[1], data.get("summary", ""))
    set_footer_line(
        document,
        str(data.get("recorder") or data.get("host") or ""),
        str(data.get("meeting_date", "")),
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    restore_static_parts(template, output)
    validate_output(output)


def main() -> None:
    args = parse_args()
    with args.input.open("r", encoding="utf-8") as file:
        data = json.load(file)
    generate(data, args.template, args.output)
    print(args.output.resolve())


if __name__ == "__main__":
    main()
